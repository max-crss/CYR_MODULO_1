"""
Generador del documento Word de ejercicios para la Actividad N°6: Instalando Puntos de Red
Módulo 1 – Instalación de redes de área local cableadas e inalámbricas
Conectividad y Redes – 3° Medio TP

Uso:
    pip install python-docx
    python generate_act6_word.py

Genera: ACT6_Ejercicios_Puntos_de_Red.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores ────────────────────────────────────────────────────────────────────
AZUL_OSCURO = RGBColor(0x1F, 0x38, 0x64)   # headers
NARANJA     = RGBColor(0xFF, 0x66, 0x00)   # acentos
BLANCO      = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_CLARO  = RGBColor(0xD4, 0xDC, 0xF0)  # filas alternas tabla


# ══════════════════════════════════════════════════════════════════════════════
# Helpers de formato
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, rgb: RGBColor):
    """Establece el color de fondo de una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, color=None, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Escribe texto con formato en una celda."""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color


def heading(doc, text, level=1):
    """Agrega un encabezado con estilo personalizado."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = AZUL_OSCURO
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = NARANJA
    return para


def answer_lines(doc, count=3, indent=False):
    """Agrega líneas de respuesta (guiones bajos)."""
    for _ in range(count):
        para = doc.add_paragraph()
        run = para.add_run('_' * 80)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if indent:
            para.paragraph_format.left_indent = Inches(0.3)
    return


def add_exercise_header(doc, number, title):
    """Encabezado de ejercicio con número en naranja y título en azul."""
    para = doc.add_paragraph()
    run_num = para.add_run(f"Ejercicio {number}: ")
    run_num.bold = True
    run_num.font.name = 'Calibri'
    run_num.font.size = Pt(13)
    run_num.font.color.rgb = NARANJA

    run_title = para.add_run(title)
    run_title.bold = True
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(13)
    run_title.font.color.rgb = AZUL_OSCURO
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    return para


def add_instruction(doc, text):
    """Instrucción del ejercicio en itálica."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_body(doc, text, bold=False):
    """Texto de cuerpo normal."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return para


def divider(doc):
    """Línea divisora entre ejercicios."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# Documento principal
# ══════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── Configuración de página: Carta, márgenes normales ─────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # Fuente predeterminada
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Encabezado del documento ───────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('ACTIVIDAD N°6 – INSTALANDO PUNTOS DE RED')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL_OSCURO

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_para.add_run('Hoja de Ejercicios Teóricos  |  Conectividad y Redes  |  3° Medio TP')
    run2.font.name = 'Calibri'
    run2.font.size = Pt(11)
    run2.font.color.rgb = NARANJA
    run2.bold = True

    doc.add_paragraph()

    # ── Datos del alumno ───────────────────────────────────────────────────────
    info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Table Grid'
    labels = ['Nombre:', 'Fecha:', 'Curso:', 'Puntaje:']
    for i, label in enumerate(labels):
        cell = info_table.rows[0].cells[i]
        cell_text(cell, label, bold=True, color=AZUL_OSCURO, size=10)
        set_cell_bg(cell, GRIS_CLARO)
        info_table.rows[1].cells[i].paragraphs[0].add_run('').font.size = Pt(11)

    # Ancho columnas datos alumno
    col_widths = [Inches(2.0), Inches(1.2), Inches(1.0), Inches(1.0)]
    for i, col in enumerate(info_table.columns):
        for cell in col.cells:
            cell.width = col_widths[i]

    doc.add_paragraph()

    # Instrucciones generales
    gen_para = doc.add_paragraph()
    run_g = gen_para.add_run(
        'Instrucciones generales: '
    )
    run_g.bold = True
    run_g.font.name = 'Calibri'
    run_g.font.size = Pt(11)
    run_g.font.color.rgb = AZUL_OSCURO
    run_g2 = gen_para.add_run(
        'Responde cada ejercicio en el espacio indicado. No se requieren materiales físicos. '
        'Puedes apoyarte en los apuntes de clases y la norma TIA 568B vista en la actividad.'
    )
    run_g2.font.name = 'Calibri'
    run_g2.font.size = Pt(11)

    divider(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # Ejercicio 1 – Unir con flechas (tabla de dos columnas)
    # ══════════════════════════════════════════════════════════════════════════
    add_exercise_header(doc, 1, 'Une cada componente con su descripción correcta')
    add_instruction(doc, 'Escribe la letra de la columna derecha que corresponda al componente de la columna izquierda.')

    match_table = doc.add_table(rows=9, cols=3)
    match_table.style = 'Table Grid'
    match_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    headers = ['Componente', '', 'Descripción']
    for j, h in enumerate(headers):
        cell = match_table.rows[0].cells[j]
        set_cell_bg(cell, AZUL_OSCURO)
        cell_text(cell, h, bold=True, color=BLANCO, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Contenido: componentes (izq) y descripciones mezcladas (der)
    componentes = [
        'Jack Keystone RJ-45',
        'Herramienta Punch-Down',
        'Patch Panel',
        'Cable UTP Cat6',
        'Faceplate',
        'LAN Tester',
        'Canaleta plástica',
        'Rack de comunicaciones',
    ]
    descripciones = [
        ('A', 'Permite verificar la continuidad y el orden correcto de los pares de un cable.'),
        ('B', 'Armario metálico que organiza todos los equipos y cables de la red.'),
        ('C', 'Toma de red que se instala en la pared; recibe los conductores del cable UTP.'),
        ('D', 'Herramienta que presiona cada conductor en su ranura del jack o patch panel.'),
        ('E', 'Soporte plástico frontal que cubre el jack y da terminación estética al punto.'),
        ('F', 'Panel centralizado con múltiples puertos RJ-45 que recibe todos los cables del tendido.'),
        ('G', 'Cable de 4 pares trenzados usado para cableado horizontal de hasta 100 m.'),
        ('H', 'Conducto plástico que protege y ordena los cables a lo largo de la pared.'),
    ]

    for i, (comp, (letra, desc)) in enumerate(zip(componentes, descripciones)):
        row = match_table.rows[i + 1]
        cell_text(row.cells[0], comp, size=11)
        cell_text(row.cells[1], '______', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[2], f'{letra}. {desc}', size=11)
        if i % 2 == 0:
            set_cell_bg(row.cells[0], GRIS_CLARO)
            set_cell_bg(row.cells[1], GRIS_CLARO)
            set_cell_bg(row.cells[2], GRIS_CLARO)

    # Ancho columnas ejercicio 1
    widths_ex1 = [Inches(2.2), Inches(0.6), Inches(3.6)]
    for j, col in enumerate(match_table.columns):
        for cell in col.cells:
            cell.width = widths_ex1[j]

    divider(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # Ejercicio 2 – Ordenar pasos
    # ══════════════════════════════════════════════════════════════════════════
    add_exercise_header(doc, 2, 'Ordena los pasos de instalación de un jack Keystone')
    add_instruction(doc,
        'Los siguientes pasos están en desorden. Escribe el número del 1 al 7 '
        'indicando el orden correcto de instalación.')

    steps_table = doc.add_table(rows=8, cols=2)
    steps_table.style = 'Table Grid'
    steps_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    set_cell_bg(steps_table.rows[0].cells[0], AZUL_OSCURO)
    set_cell_bg(steps_table.rows[0].cells[1], AZUL_OSCURO)
    cell_text(steps_table.rows[0].cells[0], 'Orden (1–7)', bold=True, color=BLANCO, size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(steps_table.rows[0].cells[1], 'Paso a realizar', bold=True, color=BLANCO, size=11)

    pasos_mezclados = [
        'Probar con LAN Tester: los 8 LEDs deben encenderse en orden del 1 al 8.',
        'Pelar 3 cm de la cubierta exterior del cable UTP con el pelacable.',
        'Fijar el faceplate a la caja empotrada en la pared.',
        'Usar la herramienta punch-down para presionar cada conductor en su ranura del jack.',
        'Marcar y cortar el cable dejando 30 cm de margen en la caja.',
        'Ordenar los 8 conductores según el diagrama de colores de la norma TIA 568B.',
        'Insertar el jack en el faceplate hasta escuchar un clic de encaje.',
    ]

    for i, paso in enumerate(pasos_mezclados):
        row = steps_table.rows[i + 1]
        cell_text(row.cells[0], '______', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[1], paso, size=11)
        if i % 2 == 0:
            set_cell_bg(row.cells[0], GRIS_CLARO)
            set_cell_bg(row.cells[1], GRIS_CLARO)

    for col_idx, width in enumerate([Inches(1.2), Inches(5.2)]):
        for cell in steps_table.columns[col_idx].cells:
            cell.width = width

    divider(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # Ejercicio 3 – Tabla de colores TIA 568B
    # ══════════════════════════════════════════════════════════════════════════
    add_exercise_header(doc, 3, 'Completa la tabla de colores TIA 568B')
    add_instruction(doc,
        'Completa los campos vacíos con el color correcto del conductor '
        'según la norma TIA 568B. Algunos pines ya están completados como referencia.')

    color_table = doc.add_table(rows=9, cols=3)
    color_table.style = 'Table Grid'
    color_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    for j, h in enumerate(['Pin', 'Color del conductor', 'Par']):
        cell = color_table.rows[0].cells[j]
        set_cell_bg(cell, NARANJA)
        cell_text(cell, h, bold=True, color=BLANCO, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Datos: pin, color (None = vacío para completar), par
    tia568b = [
        ('1', 'Blanco-Naranja',  'Par 2'),
        ('2', None,              'Par 2'),
        ('3', 'Blanco-Verde',    'Par 3'),
        ('4', None,              'Par 1'),
        ('5', 'Blanco-Azul',     'Par 1'),
        ('6', None,              'Par 3'),
        ('7', 'Blanco-Café',     'Par 4'),
        ('8', None,              'Par 4'),
    ]

    for i, (pin, color, par) in enumerate(tia568b):
        row = color_table.rows[i + 1]
        cell_text(row.cells[0], pin, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        if color:
            cell_text(row.cells[1], color, size=11)
            set_cell_bg(row.cells[1], GRIS_CLARO)
        else:
            cell_text(row.cells[1], '____________________', size=11)
        cell_text(row.cells[2], par, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    for col_idx, width in enumerate([Inches(0.6), Inches(3.4), Inches(1.2)]):
        for cell in color_table.columns[col_idx].cells:
            cell.width = width

    divider(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # Ejercicio 4 – Preguntas de desarrollo corto
    # ══════════════════════════════════════════════════════════════════════════
    add_exercise_header(doc, 4, 'Responde brevemente las siguientes preguntas')
    add_instruction(doc,
        'Responde cada pregunta con 2 o 3 líneas explicando cómo funcionan '
        'las herramientas y conexiones del punto de red.')

    preguntas = [
        ('4.1', '¿Por qué la herramienta punch-down corta el exceso de conductor al mismo tiempo que lo presiona en la ranura?'),
        ('4.2', '¿Qué sucedería con la señal de red si los conductores dentro del jack quedan destrenzados más de 1,25 cm?'),
        ('4.3', '¿Cuál es la diferencia entre un patch cord y el cable de tendido horizontal? ¿Para qué se usa cada uno?'),
        ('4.4', '¿Por qué el LAN Tester verifica los 8 pines en orden y no solo que haya continuidad eléctrica?'),
        ('4.5', '¿Qué ventaja entrega el patch panel frente a conectar los cables directamente al switch?'),
    ]

    for num, pregunta in preguntas:
        q_para = doc.add_paragraph()
        run_n = q_para.add_run(f'{num}. ')
        run_n.bold = True
        run_n.font.name = 'Calibri'
        run_n.font.size = Pt(11)
        run_n.font.color.rgb = NARANJA
        run_q = q_para.add_run(pregunta)
        run_q.font.name = 'Calibri'
        run_q.font.size = Pt(11)
        q_para.paragraph_format.space_before = Pt(6)
        answer_lines(doc, count=3, indent=True)

    divider(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # Ejercicio 5 – Verdadero o Falso
    # ══════════════════════════════════════════════════════════════════════════
    add_exercise_header(doc, 5, 'Verdadero o Falso – Justifica tu respuesta')
    add_instruction(doc,
        'Indica si cada afirmación es Verdadera (V) o Falsa (F) marcando con una X. '
        'Luego escribe una breve justificación.')

    vf_table = doc.add_table(rows=8, cols=4)
    vf_table.style = 'Table Grid'
    vf_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    for j, h in enumerate(['N°', 'Afirmación', 'V / F', 'Justificación']):
        cell = vf_table.rows[0].cells[j]
        set_cell_bg(cell, AZUL_OSCURO)
        cell_text(cell, h, bold=True, color=BLANCO, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    afirmaciones = [
        'El cable UTP Cat6 puede transmitir datos hasta 200 metros sin necesidad de repetidores.',
        'La norma TIA 568B define el orden de colores de los conductores al conectar el jack Keystone.',
        'La herramienta punch-down puede usarse tanto en jacks como en patch panels.',
        'El faceplate cumple únicamente una función estética sin importancia para el rendimiento de la red.',
        'Un LAN Tester con todos sus LEDs en orden (1-2-3-4-5-6-7-8) indica que el cable está correctamente instalado.',
        'Se puede destrensar cualquier cantidad de cable antes del punch-down sin afectar la calidad de la señal.',
        'El patch panel permite reorganizar y reubicar conexiones de red sin mover el cableado fijo de los muros.',
    ]

    for i, afirm in enumerate(afirmaciones):
        row = vf_table.rows[i + 1]
        cell_text(row.cells[0], str(i + 1), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[1], afirm, size=10)
        cell_text(row.cells[2], 'V  ☐\nF  ☐', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[3], '', size=10)
        if i % 2 == 0:
            set_cell_bg(row.cells[0], GRIS_CLARO)
            set_cell_bg(row.cells[1], GRIS_CLARO)
            set_cell_bg(row.cells[2], GRIS_CLARO)
            set_cell_bg(row.cells[3], GRIS_CLARO)

    for col_idx, width in enumerate([Inches(0.4), Inches(2.8), Inches(0.7), Inches(2.5)]):
        for cell in vf_table.columns[col_idx].cells:
            cell.width = width

    divider(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # Ejercicio 6 – Análisis de diagrama de red
    # ══════════════════════════════════════════════════════════════════════════
    add_exercise_header(doc, 6, 'Análisis del trayecto de red')
    add_instruction(doc,
        'Observa el siguiente diagrama del trayecto que recorre la señal de datos '
        'desde un computador hasta el switch, y responde las preguntas.')

    # Diagrama textual
    diag_para = doc.add_paragraph()
    diag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = diag_para.add_run(
        'PC  ──►  Jack de pared  ──►  Cable UTP en canaleta  ──►  Patch Panel  ──►  Patch Cord  ──►  Switch'
    )
    run_d.bold = True
    run_d.font.name = 'Calibri'
    run_d.font.size = Pt(12)
    run_d.font.color.rgb = AZUL_OSCURO
    diag_para.paragraph_format.space_before = Pt(6)
    diag_para.paragraph_format.space_after = Pt(10)

    preguntas_ex6 = [
        ('6.1', '¿Qué elemento del trayecto se instala empotrado en la pared y recibe directamente el cable UTP proveniente del tendido?'),
        ('6.2', '¿Cuál es la función del patch cord en este trayecto y en qué dos puntos se usa habitualmente?'),
        ('6.3', 'Si la señal no llegara al switch pero el LED del LAN Tester sí se enciende en el jack, ¿en qué tramo del trayecto sospecharías el problema?'),
        ('6.4', '¿Por qué la canaleta es importante dentro de este trayecto? Menciona al menos dos razones.'),
    ]

    for num, pregunta in preguntas_ex6:
        q_para = doc.add_paragraph()
        run_n = q_para.add_run(f'{num}. ')
        run_n.bold = True
        run_n.font.name = 'Calibri'
        run_n.font.size = Pt(11)
        run_n.font.color.rgb = NARANJA
        run_q = q_para.add_run(pregunta)
        run_q.font.name = 'Calibri'
        run_q.font.size = Pt(11)
        q_para.paragraph_format.space_before = Pt(6)
        answer_lines(doc, count=3, indent=True)

    divider(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # Ejercicio 7 – Estudio de caso
    # ══════════════════════════════════════════════════════════════════════════
    add_exercise_header(doc, 7, 'Estudio de caso – Diagnóstico de errores')
    add_instruction(doc,
        'Lee atentamente el siguiente caso y responde las preguntas identificando '
        'los errores cometidos por el técnico y cómo deberían corregirse.')

    # Texto del caso
    caso_para = doc.add_paragraph()
    run_caso = caso_para.add_run(
        'Situación: '
    )
    run_caso.bold = True
    run_caso.font.name = 'Calibri'
    run_caso.font.size = Pt(11)
    run_caso.font.color.rgb = AZUL_OSCURO

    run_caso2 = caso_para.add_run(
        'El técnico Rodrigo instaló 10 puntos de red en una oficina. Al finalizar, '
        'conectó los computadores pero ninguno tenía acceso a la red. Al revisar, notó lo siguiente:\n'
        '  • Destrenó todos los conductores hasta el conector antes de presionarlos en el jack.\n'
        '  • Usó cable UTP Cat5 que encontró en bodega, en lugar del Cat6 especificado.\n'
        '  • No etiquetó ningún cable ni en el jack ni en el patch panel.\n'
        '  • Al probar con el LAN Tester, los LEDs se encendían en el siguiente orden: 1-2-4-3-5-6-8-7.\n'
        '  • Dejó los cables sin canaleta en varios tramos del tendido.'
    )
    run_caso2.font.name = 'Calibri'
    run_caso2.font.size = Pt(11)
    caso_para.paragraph_format.space_after = Pt(8)

    preguntas_ex7 = [
        ('7.1', '¿Cuántos errores puedes identificar en la instalación de Rodrigo? Enuméralos.'),
        ('7.2', '¿Qué indica el orden 1-2-4-3-5-6-8-7 en el LAN Tester? ¿Cómo se corrige?'),
        ('7.3', '¿Por qué el uso de cable Cat5 en lugar de Cat6 podría ser un problema a largo plazo?'),
        ('7.4', '¿Qué consecuencias operativas tiene no etiquetar los cables en un patch panel de 24 puertos?'),
    ]

    for num, pregunta in preguntas_ex7:
        q_para = doc.add_paragraph()
        run_n = q_para.add_run(f'{num}. ')
        run_n.bold = True
        run_n.font.name = 'Calibri'
        run_n.font.size = Pt(11)
        run_n.font.color.rgb = NARANJA
        run_q = q_para.add_run(pregunta)
        run_q.font.name = 'Calibri'
        run_q.font.size = Pt(11)
        q_para.paragraph_format.space_before = Pt(6)
        answer_lines(doc, count=3, indent=True)

    # ── Pie de página ──────────────────────────────────────────────────────────
    divider(doc)
    foot_para = doc.add_paragraph()
    foot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = foot_para.add_run(
        'Actividad N°6  |  Instalando Puntos de Red  |  Conectividad y Redes – 3° Medio TP  |  '
        'Total: ____ / ____ pts'
    )
    run_f.font.name = 'Calibri'
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = AZUL_OSCURO

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# Entry point
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    output = 'ACT6_Ejercicios_Puntos_de_Red.docx'
    doc = build_document()
    doc.save(output)
    print(f'✅  Documento generado: {output}')
